
import os
import io
import re
import json
import unicodedata
from typing import List, Dict, Any, Tuple

import numpy as np
import pandas as pd
from fastapi import FastAPI, HTTPException, Header, Response, Query
from fastapi.responses import StreamingResponse, JSONResponse
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from pdfminer.high_level import extract_text as pdf_extract_text

from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2 import service_account

# =========================
#  Seguridad por API Key
# =========================
API_KEY = os.environ.get("API_KEY", "")

app = FastAPI(title="Generador CES")

# =========================
#  Autenticación Google Drive (Service Account)
# =========================
SCOPES = [s.strip() for s in os.environ.get("GOOGLE_SCOPES", "https://www.googleapis.com/auth/drive.readonly").split(",")]

def get_drive_service():
    sa_str = os.environ.get("GCP_SA_KEY")
    if not sa_str:
        raise HTTPException(status_code=500, detail="Falta GCP_SA_KEY en variables de entorno")
    try:
        info = json.loads(sa_str)
    except Exception:
        raise HTTPException(status_code=500, detail="GCP_SA_KEY inválida (no es JSON)")
    creds = service_account.Credentials.from_service_account_info(info, scopes=SCOPES)
    return build('drive', 'v3', credentials=creds)

# =========================
#  Utilidades y Normalización
# =========================
KV_PATTERN = re.compile(r"\b\d+\s*[kK]\s*[vV]\b")

TRANSLATION_TRAMOS = str.maketrans({
    "á":"a","Á":"A","é":"e","É":"E","í":"i","Í":"I","ó":"o","Ó":"O","ú":"u","Ú":"U","ü":"u","Ü":"U"
})

def strip_accents(s: str) -> str:
    s = unicodedata.normalize("NFD", str(s))
    return "".join(ch for ch in s if unicodedata.category(ch) != "Mn")

def strip_accents_and_lower(s: str) -> str:
    return strip_accents(s).lower()

def normalize_dashes(s: str) -> str:
    return str(s).replace("–", "-").replace("—", "-").replace("−", "-")

def norm_text_local(s: str) -> str:
    s = unicodedata.normalize("NFD", str(s))
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = s.lower()
    s = re.sub(r"\s+", " ", s).strip()
    return s

def norm_text(s: str) -> str:
    if not isinstance(s, str):
        s = "" if pd.isna(s) else str(s)
    s = s.replace("–", "-")
    s = strip_accents(s).lower().strip()
    s = re.sub(r"\s+", " ", s)
    return s

# =========================
#  Drive helpers
# =========================
def extract_folder_id(url: str) -> str:
    m = re.search(r"/folders/([a-zA-Z0-9_-]+)", url)
    if not m:
        raise HTTPException(status_code=400, detail="No se pudo extraer el ID de la carpeta desde FOLDER_URL")
    return m.group(1)

def _find_by_name_contains_with_mimes(drive_service, folder_id: str, name_contains: str, mimes: list, label: str):
    mime_filter = " or ".join([f"mimeType = '{m}'" for m in mimes])
    q = (f"'{folder_id}' in parents and trashed = false and "
         f"name contains '{name_contains}' and (" + mime_filter + ")")
    resp = drive_service.files().list(q=q, spaces='drive', includeItemsFromAllDrives=True,
                                      supportsAllDrives=True,
                                      fields="files(id,name,mimeType,modifiedTime)",
                                      orderBy="modifiedTime desc", pageSize=100).execute()
    files = resp.get('files', [])
    if files:
        return files[0]
    q_name_only = f"'{folder_id}' in parents and trashed = false and name contains '{name_contains}'"
    resp2 = drive_service.files().list(q=q_name_only, spaces='drive', includeItemsFromAllDrives=True,
                                       supportsAllDrives=True,
                                       fields="files(id,name,mimeType,modifiedTime)",
                                       orderBy="modifiedTime desc", pageSize=100).execute()
    found = resp2.get('files', [])
    if found:
        return found[0]
    raise HTTPException(status_code=404, detail=f"No se encontró archivo que contenga '{name_contains}'.")

def find_excel_like_in_folder(drive_service, folder_id: str, name_contains: str):
    mimes = [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel.sheet.macroEnabled.12",
        "application/vnd.google-apps.spreadsheet",
    ]
    return _find_by_name_contains_with_mimes(drive_service, folder_id, name_contains, mimes, "FTS")

def find_exact_xlsx_in_folder(drive_service, folder_id: str, filename: str):
    q = (f"'{folder_id}' in parents and trashed = false and "
         f"name = '{filename}' and mimeType = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'")
    resp = drive_service.files().list(q=q, spaces='drive', includeItemsFromAllDrives=True, supportsAllDrives=True,
                                      fields="files(id,name,mimeType,modifiedTime)", orderBy="modifiedTime desc", pageSize=20).execute()
    files = resp.get("files", [])
    if files:
        return files[0]
    base = filename.rsplit(".", 1)[0]
    q_like = f"'{folder_id}' in parents and trashed = false and name contains '{base}'"
    resp2 = drive_service.files().list(q=q_like, spaces='drive', includeItemsFromAllDrives=True, supportsAllDrives=True,
                                       fields="files(id,name,mimeType,modifiedTime)", orderBy="modifiedTime desc", pageSize=50).execute()
    found = resp2.get("files", [])
    if found:
        return found[0]
    raise HTTPException(status_code=404, detail=f"No se encontró '{filename}'.")

def download_to_xlsx(drive_service, file_meta: dict, dest_io: io.BytesIO) -> bytes:
    file_id = file_meta["id"]
    mime = file_meta["mimeType"]
    if mime == "application/vnd.google-apps.spreadsheet":
        request = drive_service.files().export_media(fileId=file_id, mimeType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        request = drive_service.files().get_media(fileId=file_id)
    downloader = MediaIoBaseDownload(dest_io, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    dest_io.seek(0)
    return dest_io.getvalue()

def download_to_docx(drive_service, file_meta: dict, dest_io: io.BytesIO) -> bytes:
    file_id = file_meta["id"]
    mime = file_meta["mimeType"]
    if mime == "application/vnd.google-apps.document":
        request = drive_service.files().export_media(fileId=file_id, mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        request = drive_service.files().get_media(fileId=file_id)
    downloader = MediaIoBaseDownload(dest_io, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    dest_io.seek(0)
    return dest_io.getvalue()

def list_pdfs_in_folder(drive_service, folder_id: str, page_size: int = 200):
    q = (f"'{folder_id}' in parents and trashed = false and mimeType = 'application/pdf'")
    resp = drive_service.files().list(q=q, spaces='drive', includeItemsFromAllDrives=True, supportsAllDrives=True,
                                      fields="files(id,name,mimeType,modifiedTime)", orderBy="modifiedTime desc", pageSize=page_size).execute()
    return resp.get("files", [])

# =========================
#  Lógica (núcleo mínimo)
# =========================
def read_nombre_linea_asociada_column(xlsx_bytes: bytes, sheet_name: str) -> list:
    import openpyxl
    bio = io.BytesIO(xlsx_bytes)
    df = pd.read_excel(bio, sheet_name=sheet_name, header=None, dtype=str, engine="openpyxl")
    df = df.fillna("")
    target_header_norm = norm_text("Nombre Línea Asociada")
    header_pos = None
    for r in range(df.shape[0]):
        for c in range(df.shape[1]):
            val = str(df.iat[r, c]).strip()
            if val and norm_text(val) == target_header_norm:
                header_pos = (r, c)
                break
        if header_pos:
            break
    if not header_pos:
        raise HTTPException(status_code=400, detail="No se encontró la columna 'Nombre Línea Asociada'.")
    r0, c0 = header_pos
    col_vals = [str(v).strip() for v in df.iloc[r0+1:, c0].tolist()]
    return [v for v in col_vals if v and v.lower() != "nan"]

def parse_line_names(raw_value: str):
    if not isinstance(raw_value, str):
        return None, None, False
    txt = normalize_dashes(raw_value)
    u_txt = strip_accents_and_lower(txt)
    if "[no_mostrar]" in u_txt:
        return None, None, False
    if "-" not in txt:
        return None, None, False
    left, right = txt.split("-", 1)
    sub1 = left.rstrip().strip()
    rr = right.lstrip()
    m = KV_PATTERN.search(rr)
    sub2 = rr[:m.start()].strip() if m else rr.strip()
    if not sub1 or not sub2:
        return None, None, False
    return sub1, sub2, True

def build_graph(values: list):
    display_name = {}
    adj = {}
    def add_node(name):
        k = norm_text(name)
        if k not in display_name:
            display_name[k] = name.strip()
        if k not in adj:
            adj[k] = set()
        return k
    for raw in values:
        sub1, sub2, ok = parse_line_names(raw)
        if not ok:
            continue
        k1 = add_node(sub1)
        k2 = add_node(sub2)
        adj[k1].add(k2)
        adj[k2].add(k1)
    return display_name, adj

def compute_hierarchy_and_targets(display_name, adj, query: str):
    qk = norm_text(query)
    if qk not in adj or not adj[qk]:
        txt = f"Ingreso: {query}\nRespuesta:\n- {query}"
        return txt, [query.strip()]
    lvl1_items = [(display_name.get(nk, nk), nk) for nk in adj[qk]]
    lvl1_items.sort(key=lambda x: strip_accents_and_lower(x[0]))
    lines = []
    lines.append(f"Ingreso: {query}")
    lines.append("Respuesta:")
    lines.append(f" {query}")
    targets = []
    def push_unique(name):
        if name not in targets:
            targets.append(name)
    push_unique(display_name.get(qk, query.strip()))
    for disp1, key1 in lvl1_items:
        lines.append(f" {disp1}")
        push_unique(disp1)
        lvl2 = [(display_name.get(nk2, nk2), nk2) for nk2 in adj[key1] if nk2 != qk]
        dedup = {k2: dname for dname, k2 in lvl2}
        lvl2_unique = [(v, k) for k, v in dedup.items()]
        lvl2_unique.sort(key=lambda x: strip_accents_and_lower(x[0]))
        for dname2, _ in lvl2_unique:
            lines.append(f"   {dname2}")
            push_unique(dname2)
    return "\n".join(lines), targets

def find_template_docx_in_folder(drive_service, folder_id: str, name_contains: str):
    mimes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.google-apps.document",
    ]
    return _find_by_name_contains_with_mimes(drive_service, folder_id, name_contains, mimes, "DOCX")

def find_pdf_resolucion_exenta(drive_service, folder_id: str, keywords=("resolucion","exenta")):
    files = list_pdfs_in_folder(drive_service, folder_id)
    kws = [norm_text_local(k) for k in keywords]
    for f in files:
        name_norm = norm_text_local(f["name"])
        if all(kw in name_norm for kw in kws):
            return f
    if files:
        return files[0]
    raise HTTPException(status_code=404, detail="No hay PDFs en la carpeta")

def extract_resolucion_info_from_pdf(pdf_bytes: bytes):
    try:
        text = pdf_extract_text(io.BytesIO(pdf_bytes)) or ""
    except Exception:
        return "", ""
    text_norm = norm_text_local(text)
    mnum = re.search(r"resolucion\s+exenta\s+n[º°o]?\s*([0-9]+)", text_norm)
    numero = mnum.group(1) if mnum else ""
    mfecha = re.search(r"SANTIAGO,\s*([^\n\r]+)", text, flags=re.IGNORECASE)
    fecha = mfecha.group(1).strip() if mfecha else ""
    return numero, fecha

def insert_table_after(paragraph, rows, cols):
    new_tbl = OxmlElement("w:tbl")
    paragraph._p.addnext(new_tbl)
    from docx.table import Table
    tbl = Table(new_tbl, paragraph._parent)
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(cols):
        gridCol = OxmlElement('w:gridCol')
        tblGrid.append(gridCol)
    tbl._tbl.append(tblGrid)
    for _ in range(rows):
        tr = OxmlElement('w:tr')
        for _c in range(cols):
            tc = OxmlElement('w:tc')
            p = OxmlElement('w:p')
            tc.append(p)
            tr.append(tc)
        tbl._tbl.append(tr)
    return tbl

def fill_table_from_df(tbl, df: pd.DataFrame):
    hdr = df.columns.tolist()
    for j, text in enumerate(hdr):
        cell = tbl.rows[0].cells[j]
        cell.text = str(text)
        if cell.paragraphs and cell.paragraphs[0].runs:
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            val = "" if pd.isna(df.iat[i, j]) else str(df.iat[i, j])
            tbl.rows[i+1].cells[j].text = val

def replace_marker_with_table(doc: Document, markers: list, df: pd.DataFrame):
    if df is None or df.empty:
        df = pd.DataFrame({"Sin datos": ["(No se encontraron coincidencias)"]})
    paragraphs = list(doc.paragraphs)
    for p in paragraphs:
        full_text = p.text.strip()
        if any(mk in full_text for mk in markers):
            rows = 1 + len(df)
            cols = len(df.columns)
            tbl = insert_table_after(p, rows, cols)
            fill_table_from_df(tbl, df)
            for run in p.runs:
                run.text = ""
            if not p.runs:
                p.text = ""

def replace_text_everywhere(doc: Document, replacements: dict):
    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                txt = run.text
                for k, v in replacements.items():
                    if k in txt:
                        run.text = txt.replace(k, v)
    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)
    for section in doc.sections:
        for hdrftr in (section.header, section.footer):
            replace_in_paragraphs(hdrftr.paragraphs)
            for table in hdrftr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_paragraphs(cell.paragraphs)

def filename_from_subestacion(name: str) -> str:
    if not name:
        return "CES-Documento"
    s = str(name)
    s = normalize_dashes(s)
    s = re.sub(r"\s+", " ", s).strip()
    s = s.title()
    s = re.sub(r'[\\/*?:"<>|]', '-', s)
    s = s.strip(" .")
    return f"CES-{s}"

def pretty_subestacion(name: str) -> str:
    if not name:
        return ""
    s = normalize_dashes(str(name))
    s = re.sub(r"\s+", " ", s).strip()
    return s.title()

# =========================
#  Endpoints
# =========================
@app.get("/health")
def health():
    return {"ok": True}

@app.post("/generate")
async def generate(
    subestacion: str = Query(..., description="Nombre de subestación"),
    folder_url: str = Query(..., description="URL de carpeta de Google Drive"),
    x_api_key: str = Header(default="")
):
    if API_KEY and x_api_key != API_KEY:
        raise HTTPException(status_code=401, detail="API key inválida")

    drive = get_drive_service()
    folder_id = extract_folder_id(folder_url)

    # 1) Adyacencias (mínimo viable)
    fts_meta = find_excel_like_in_folder(drive, folder_id, "FichaTecnicaSecciones")
    fts_bytes = download_to_xlsx(drive, fts_meta, io.BytesIO())
    valores = read_nombre_linea_asociada_column(fts_bytes, "Patron")
    display_name, adj = build_graph(valores)
    texto_adyacencias, targets = compute_hierarchy_and_targets(display_name, adj, subestacion)

    # 2) Placeholders de tablas (puedes ampliar luego)
    df_decl = pd.DataFrame(columns=["Valor de entrada","Valor de salida","Hoja","Marca"])
    df_proy = pd.DataFrame(columns=["Término Buscado","Hoja","Contenido Coincidente","Nombre Proyecto (si aplica)","NUP (sin coma ni decimales)","Fecha Real de EO (si aplica)"])
    df_sol  = pd.DataFrame(columns=["NUP","Proyecto","Tipo Proyecto","Punto de Conexión","Capacidad [MW]"])
    df_cont = pd.DataFrame(columns=["Término","Contenido"])
    df_trafo= pd.DataFrame(columns=["Archivo","Hoja","Columna A","Columna D"])

    # 3) Resolución exenta (número y fecha)
    try:
        pdf_meta = find_pdf_resolucion_exenta(drive, folder_id)
        buf_pdf = io.BytesIO()
        request = drive.files().get_media(fileId=pdf_meta["id"])
        downloader = MediaIoBaseDownload(buf_pdf, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        buf_pdf.seek(0)
        numero_res, fecha_res = extract_resolucion_info_from_pdf(buf_pdf.getvalue())
    except Exception:
        numero_res, fecha_res = "", ""

    # 4) Plantilla DOCX
    tpl_meta = find_template_docx_in_folder(drive, folder_id, "YYMM-DEN-CES-PR")
    tpl_io = io.BytesIO()
    download_to_docx(drive, tpl_meta, tpl_io)

    doc = Document(tpl_io)
    replace_marker_with_table(doc, ["LISTADECRETADOS"], df_decl)
    replace_marker_with_table(doc, ["LISTACONEXIONES", "LISTA CONEXIONES"], df_proy)
    replace_marker_with_table(doc, ["TABLACES"], df_sol)
    replace_marker_with_table(doc, ["LISTACONTINGENCIAS"], df_cont)
    replace_marker_with_table(doc, ["LISTATRANSFORMADORES"], df_trafo)

    replacements = {
        "NOMBRESUBESTACION": pretty_subestacion(subestacion),
        "NUMERORESOLUCION":  numero_res or "",
        "FECHARESOLUCION":   fecha_res or "",
    }
    replace_text_everywhere(doc, replacements)

    out_name = filename_from_subestacion(subestacion) + ".docx"
    out_io = io.BytesIO()
    doc.save(out_io)
    out_io.seek(0)

    headers = {'Content-Disposition': f'attachment; filename="{out_name}"'}
    return StreamingResponse(out_io, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)
